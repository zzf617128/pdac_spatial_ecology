from pathlib import Path
import re
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
MANUSCRIPT = Path(os.environ.get("MANUSCRIPT_MD", PROJECT / "results" / "manuscript" / "pdac_caf_myeloid_spatial_niche_submission_v2_full.md"))
OUT = Path(os.environ.get("OUT_DOCX", PROJECT / "results" / "manuscript" / "pdac_caf_myeloid_spatial_niche_submission_v2_full.docx"))
NATURE_STYLE = os.environ.get("NATURE_STYLE", "0") == "1"


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def shade_cell(cell, fill: str = "F4F6F9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_runs_from_markdown(paragraph, text: str) -> None:
    # Minimal inline parser for manuscript emphasis; leaves citations/paths unchanged.
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 2.0 if NATURE_STYLE else 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = doc.styles["Normal"]
    footer.add_run("PDAC spatial ecology submission draft").font.size = Pt(9)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    widths = [6.5 / max(len(rows[0]), 1)] * max(len(rows[0]), 1)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value)
            if r_idx == 0:
                shade_cell(cell)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()


def main() -> None:
    text = MANUSCRIPT.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    pending_paragraph: list[str] = []
    pending_table: list[str] = []
    title_added = False

    def flush_paragraph() -> None:
        nonlocal pending_paragraph
        if pending_paragraph:
            paragraph_text = " ".join(pending_paragraph).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if NATURE_STYLE else WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_from_markdown(p, paragraph_text)
            pending_paragraph = []

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            add_markdown_table(doc, pending_table)
            pending_table = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            flush_table()
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            pending_table.append(stripped)
            continue

        flush_table()

        if stripped.startswith("# "):
            flush_paragraph()
            if not title_added:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run(stripped[2:].strip())
                run.font.name = "Calibri"
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
                title_added = True
            else:
                doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            flush_paragraph()
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            flush_paragraph()
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            flush_paragraph()
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_runs_from_markdown(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_runs_from_markdown(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            pending_paragraph.append(stripped)

    flush_paragraph()
    flush_table()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
