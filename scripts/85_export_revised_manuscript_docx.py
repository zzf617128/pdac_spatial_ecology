from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = PROJECT_ROOT / "results" / "revision_2026_06_29" / "manuscript" / "Manuscript_NatureSubjournal_revised.md"
OUTPUT_DOCX = PROJECT_ROOT / "results" / "revision_2026_06_29" / "manuscript" / "Manuscript_NatureSubjournal_revised.docx"


def set_paragraph_spacing(style, before: int, after: int, line_spacing: float) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    set_paragraph_spacing(normal, 0, 8, 1.333)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        set_paragraph_spacing(style, before, after, 1.15)

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        set_paragraph_spacing(style, 0, 4, 1.208)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc: Document, text: str, style: str | None = None, justify: bool = True) -> None:
    paragraph = doc.add_paragraph(style=style)
    if justify and style in (None, "Normal"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(paragraph, text)


def build_docx() -> None:
    doc = Document()
    configure_styles(doc)

    lines = INPUT_MD.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line[2:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("000000")
            paragraph.paragraph_format.space_after = Pt(10)
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
        elif line.startswith("- "):
            add_paragraph(doc, line[2:].strip(), style="List Bullet", justify=False)
        elif re.match(r"^\d+\.\s+", line):
            add_paragraph(doc, re.sub(r"^\d+\.\s+", "", line), style="List Number", justify=False)
        else:
            add_paragraph(doc, line)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
